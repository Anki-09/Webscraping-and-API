import csv
import requests
import bs4
import openpyxl
import docx
import os
import sys
import logging
from openpyxl.styles import Font,Style  #regarding Style error here... check for new version openpyxl usage
csvfile=input('Enter the csv filename containing the links:')
file=open(csvfile,'r')                        #open the csv file containing the links
print('Reading '+csvfile+'...')
read=csv.reader(file)                            #To read the demo.csv file
doc=docx.Document()                              #create empty document
list=[]
for row in read:                                 #iterate through each row in demo.csv
    for link in row:                             #iterate through each link in the row
        try:
          req = requests.get(link)               #retrieve the data
          req.raise_for_status()                 #incase there are any errors in retrieving
        except:
          list.append(link)                      #Add the links to the list whose data could'nt be retrieved
          continue
        print('Retrieving links...')
        doc.add_heading(link,0)                  #Adding the links as headings in the document
        print('Writing to document...')
        doc.add_paragraph(req.text,'Subtitle')   #retrieved data is added to the document with 'Subtitle' style
file.close()                                     #close the demo.csv file
if len(list)==0:                                 #if no links are added to the list
    print("No unsuccessful links")
else:
    file=open('demo1.csv','a')                   #else open the demo1.csv file to write the links into it,which could'nt be retrieved
    print('Writing unsuccessful links to demo1.csv...')
    csvW=csv.writer(file)                        #To write to demo1.csv file
    csvW.writerow(list)                          #write the list into demo1.csv
    print("Unsuccessful links are written to demo1.csv file")
    file.close()                                 #close demo1.csv file
file=open('demo.docx','wb')                      #create and/or open the demo.docx in write binary mode
doc.save(file)                                   #save the created document to demo.docx
file.close()                                     #close demo.docx
print("Retrieved data from the links are written to demo.docx")

logging.basicConfig(level=logging.DEBUG,format='%(asctime)s-%(levelname)s-%(message)s')
wb=openpyxl.load_workbook('flipkart.xlsx')  #load excel file -flipkart.xlsx
sheet=wb.get_sheet_by_name('Sheet1')        #get sheet named Sheet1
doc=docx.Document()                         #create empty document
for r in range(1,sheet.get_highest_row()+1):          #iterate through each row
   url=sheet['A'+str(r)].value              #get the url values from first column
   image = sheet['B'+str(r)].value          #get the image links from second column
   try:
      res = requests.get(url)               #retrieve the data
      imglink = requests.get(image)
      res.raise_for_status()                #incase there are any errors in retrieving
      imglink.raise_for_status()
   except:
      logging.error("Couldn't load product "+str(r))  #for data which couldn't be retrieved
      continue
   print('Retrieving the data....')
   file = open(os.path.join('images','image'+str(r)+'.jpeg'), 'wb')  #open or create image file inside images folder
   for chunk in imglink.iter_content(100000):
       file.write(chunk)                      #write the image retrieved to opened file
   file.close()                               #close image file
   bs=bs4.BeautifulSoup(res.text)             #pass the retrieved data to Beautiful soup
   tv = bs.select(".B_NuCI")                  #retrieve tags with class name .B_NuCI
   print('Writing to document....')
   doc.add_heading(str(r)+")"+tv[0].getText(),2)  #get the first data from retrieved list of tags and add it as heading to the document
   tv=bs.select("._30jeq3")                   #retrieve tags with class name .30jeq3
   doc.add_paragraph(tv[0].getText(),'Subtitle')  #get the first data from retrieved list of tags and add it to the document with Subtitle style
   doc.add_picture('images\image'+str(r)+'.jpeg')  #insert the image downloaded into document
doc.save('flipkart.docx')                          #save the document
print("Check TV Product details in flipkart.docx")

url1='https://www.limeroad.com/silver-metal-fossil-p16485187?imgIdx=3&src_id=searchTrendingRail__0' #url of limeroad shopping website
url2='https://www.reliancedigital.in/fossil-fb-01-hybrid-ftw1198-smart-watch-blue/p/491615755'      #url of reliance shopping website

wb=openpyxl.Workbook()   #create excel file
sheet=wb.get_sheet_by_name('Sheet')  #get sheet named Sheet(default)
sheet['A1']='Website'                #Add heading names to the first row
sheet['B1']='Product Name'
sheet['C1']='Product Price'
sheet['D1']='Product Color'
font1=Font(name='Verdana',size=20,bold=True)  #create font object with required style
style=Style(font1)                            #create style object by passing font object as parameter
sheet['A1'].style=style                     #Add styles to headings
sheet['B1'].style=style
sheet['C1'].style=style
sheet['D1'].style=style

font=Font(name='Verdana',size=15,italic=True)  #create font object with required style
style=Style(font)                              #create style object by passing font object as parameter
try:
    res=requests.get(url1)                     #retrieve the data
    res.raise_for_status()                     #incase there are any errors in retrieving
except:
    sys.exit('Invalid url1')                       #if url couldn't be retrieved
print('Retrieving the data from: ' + str(url1))
bs=bs4.BeautifulSoup(res.text)                #pass the retrieved data to Beautiful soup
sheet['A2']='Limeroad'                        #Add data to the row
sheet['A2'].style=style                       #Add styles
name1=bs.select(".ftwN")                      #retrieve tags with class name .ftwN
sheet['B2']=name1[0].getText()                #get the first data from retrieved list of tags
sheet['B2'].style=style
price1=bs.select(".sell")                     #retrieve tags with class name .sell
sheet['C2']=price1[0].getText()               #get the first data from retrieved list of tags
sheet['C2'].style=style
color1=bs.select(".dIb")                      #retrieve tags with class name .dIb
sheet['D2']=color1[52].getText()              #get the data at index 52 from retrieved list of tags
sheet['D2'].style=style

try:
    res=requests.get(url2)                   #retrieve the data
    res.raise_for_status()                   #incase there are any errors in retrieving
except:
    sys.exit('Invalid url2')                 #if url couldn't be retrieved
print('Retrieving the data from: '+str(url2))
bs=bs4.BeautifulSoup(res.text)               #pass the retrieved data to Beautiful soup
sheet['A3']='Reliance Digital'                #Add data to the row
sheet['A3'].style=style                       #Add styles
name2=bs.select(".pdp__title")                #retrieve tags with class name .pdp__title
sheet['B3']=name2[0].getText()                #get the first data from retrieved list of tags
sheet['B3'].style=style
price2=bs.select(".pdp__offerPrice")          #retrieve tags with class name .pdp__offerPrice
sheet['C3']=price2[0].getText()[1:]             #get the first data from retrieved list of tags
sheet['C3'].style=style
color2=bs.select(".pdp__tab-info__list__value")  #retrieve tags with class name .pdp__tab-info__list__value
sheet['D3']=color2[11].getText()                 #get the data at index 11 from retrieved list of tags
sheet['D3'].style=style

wb.save('scrap.xlsx')                           #save excel file

print('Comparing cost of products.....')
if sheet['C2'].value < sheet['C3'].value:      #if cost of first product is less than second product
    print(sheet['B2'].value +' from '+sheet['A2'].value+' has less price today: '+str(sheet['C2'].value))
else:
    print(sheet['B3'].value +' from '+sheet['A3'].value+' has less price today: '+str(sheet['C3'].value))
print('Check scrap.xlsx file for all product details.')